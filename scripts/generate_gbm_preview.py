import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os


S0 = 100
mu = 0.05
sigma = 0.2
steps = 252
dt = 1 / 252


prices = [S0]

for _ in range(steps):
    z = np.random.normal()
    S_prev = prices[-1]
    S_next = S_prev * np.exp((mu - 0.5 * sigma**2) * dt + sigma * np.sqrt(dt) * z)
    prices.append(S_next)


os.makedirs("data", exist_ok=True)

chart_path = "data/gbm_chart.png"

plt.figure()

plt.plot(prices)
plt.title("Synthetic GBM Price Path")
plt.xlabel("Time Step")
plt.ylabel("Price")

plt.savefig(chart_path)
plt.close()


doc = Document()

doc.add_heading("Synthetic GBM Data Preview", level=1)

doc.add_paragraph(
    "This document shows a simulated Geometric Brownian Motion price path "
    "used for training the reinforcement learning hedging model."
)


doc.add_heading("GBM Price Chart", level=2)
doc.add_picture(chart_path, width=Inches(6))


doc.add_heading("GBM Price Table", level=2)

table = doc.add_table(rows=1, cols=2)

header = table.rows[0].cells
header[0].text = "Time Step"
header[1].text = "Price"

for i, price in enumerate(prices):
    row = table.add_row().cells
    row[0].text = str(i)
    row[1].text = f"{price:.4f}"


output_file = "data/gbm_preview.docx"
doc.save(output_file)

print("GBM preview document created:")
print(output_file)